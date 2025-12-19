import streamlit as st
from chat_agent import ChatAgent
import base64
from io import BytesIO
from docx import Document
from docx.shared import Pt
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
import os, json, time
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import re

st.markdown("", unsafe_allow_html=True)

if "chat_agent" not in st.session_state:
    st.session_state.chat_agent = ChatAgent()
if "uploader_key" not in st.session_state:
    st.session_state.uploader_key = 0
if "generated_text" not in st.session_state:
    st.session_state.generated_text = ""

def parse_bold_text(text: str):
    pdf_text = re.sub(r'\*(.*?)\*', r'<b>\1</b>', text)
    word_segments = []
    last_end = 0
    for match in re.finditer(r'\*(.*?)\*', text):
        if match.start() > last_end:
            word_segments.append((text[last_end:match.start()], False))
        word_segments.append((match.group(1), True))
        last_end = match.end()
    if last_end < len(text):
        word_segments.append((text[last_end:], False))
    return pdf_text, word_segments

def export_as_pdf(text: str) -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('TitleEnhanced', parent=styles['Title'], alignment=1,
                                 textColor=colors.HexColor('#1A73E8'), fontSize=22, spaceAfter=20)
    body_style = ParagraphStyle('BodyEnhanced', parent=styles['BodyText'], fontSize=12, leading=18)
    story = []
    banner_table = Table([["CoachPro – Document Professionnel"]], colWidths=[16*cm])
    banner_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#1A73E8')),
        ('TEXTCOLOR', (0,0), (-1,-1), colors.white),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTSIZE', (0,0), (-1,-1), 14),
        ('BOTTOMPADDING', (0,0), (-1,-1),12),
        ('TOPPADDING', (0,0), (-1,-1),12)
    ]))
    story.append(banner_table)
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph("Document Généré par CoachPro", title_style))
    story.append(Spacer(1, 0.5*cm))
    for line in text.split("\n"):
        pdf_line, _ = parse_bold_text(line)
        story.append(Paragraph(pdf_line, body_style))
        story.append(Spacer(1, 0.2*cm))
    doc.build(story)
    buffer.seek(0)
    return buffer

def export_as_word(text: str) -> BytesIO:
    buffer = BytesIO()
    document = Document()
    document.add_heading("Document Généré par CoachPro", level=1)
    for line in text.split("\n"):
        _, segments = parse_bold_text(line)
        para = document.add_paragraph()
        for segment_text, is_bold in segments:
            run = para.add_run(segment_text)
            run.bold = is_bold
            run.font.size = Pt(12)
    document.save(buffer)
    buffer.seek(0)
    return buffer

def send_mail(recipient, subject, body, attachment_name, attachment_bytes):
    sender = "brunotsogbe6@gmail.com"
    password = "eycr umjz dduf kwpk"
    msg = MIMEMultipart()
    msg['From'] = sender
    msg['To'] = recipient
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain'))
    part = MIMEBase('application', 'octet-stream')
    part.set_payload(attachment_bytes)
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', f'attachment; filename={attachment_name}')
    msg.attach(part)
    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.starttls()
    server.login(sender, password)
    server.send_message(msg)
    server.quit()

def init_header():
    st.set_page_config(page_title="CoachPro", page_icon="", layout="wide")
    st.markdown("<h1 style='text-align: center;'>💼 CoachPro</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; font-size: 16px;'>Ton coach professionnel.</p>", unsafe_allow_html=True)

def show_discussion_history(history_placeholder):
    with history_placeholder.container():
        for message in st.session_state.chat_agent.history:
            if message["role"] != "system":
                with st.chat_message(message["role"]):
                    if isinstance(message["content"], str):
                        st.write(message["content"])
                    elif isinstance(message["content"], list):
                        text = message["content"][0]["text"]
                        img_b64 = message["content"][1]["image_url"]["url"]
                        st.write(text)
                        st.image(img_b64)

HISTORY_DIR = "history"
DOC_DIR = os.path.join(HISTORY_DIR, "documents")
DISCUSSION_FILE = os.path.join(HISTORY_DIR, "discussions.json")
os.makedirs(DOC_DIR, exist_ok=True)
if not os.path.exists(DISCUSSION_FILE):
    with open(DISCUSSION_FILE, "w") as f:
        json.dump([], f)

def save_discussion(role, content):
    with open(DISCUSSION_FILE, "r") as f:
        data = json.load(f)
    data.append({"timestamp": time.time(), "role": role, "content": content})
    with open(DISCUSSION_FILE, "w") as f:
        json.dump(data, f, indent=2)

def user_interface():
    init_header()
    history_placeholder = st.empty()
    show_discussion_history(history_placeholder)

    _, col_input, col_upload = st.columns([1, 4, 2])
    user_input = None
    uploaded_file = None
    with col_input:
        st.markdown("<div style='max-width: 400px;'>", unsafe_allow_html=True)
        user_input = st.chat_input("Veillez saisir votre demande ici...", key="chat_input", max_chars=300)
        st.markdown("</div>", unsafe_allow_html=True)
    with col_upload:
        uploaded_file = st.file_uploader("Fichier", type=["png","jpg","jpeg","pdf"], accept_multiple_files=False,
                                         key=f"uploader_{st.session_state.uploader_key}")

    if user_input:
        if uploaded_file:
            if uploaded_file.type == "application/pdf":
                b64_file = base64.b64encode(uploaded_file.read()).decode("utf-8")
                response = st.session_state.chat_agent.ask_vision_model(user_interaction=user_input, image_b64=b64_file)
            else:
                image_b64 = ChatAgent.format_streamlit_image_to_base64(uploaded_file)
                response = st.session_state.chat_agent.ask_vision_model(user_interaction=user_input, image_b64=image_b64)
        else:
            response = st.session_state.chat_agent.ask_llm(user_interaction=user_input)
        st.session_state.generated_text = response if isinstance(response,str) else ""
        show_discussion_history(history_placeholder)
        st.session_state.uploader_key += 1

    if st.session_state.generated_text:
        st.markdown("Exporter le résultat")
        pdf_buffer = export_as_pdf(st.session_state.generated_text)
        word_buffer = export_as_word(st.session_state.generated_text)
        st.download_button("Télécharger en PDF", pdf_buffer, "document.pdf", "application/pdf")
        st.download_button("Télécharger en Word", word_buffer, "document.docx",
                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        st.markdown("Envoyer par mail")
        _, col_input, _ = st.columns([1,4,1])
        with col_input:
            recipient = st.text_input("Adresse e-mail du destinataire", max_chars=100)
            subject = st.text_input("Sujet", "Votre document généré par CoachPro", max_chars=100)
            body = st.text_area("Message", "Bonjour,\nVeuillez trouver ci-joint votre document généré par CoachPro.", height=120)

            if st.button("Envoyer le PDF"):
                try:
                    send_mail(recipient, subject, body, "document.pdf", pdf_buffer.getvalue())
                    st.success("PDF envoyé !")
                except Exception as e:
                    st.error(f"Erreur : {e}")
            if st.button("Envoyer le Word"):
                try:
                    send_mail(recipient, subject, body, "document.docx", word_buffer.getvalue())
                    st.success("Word envoyé !")
                except Exception as e:
                    st.error(f"Erreur : {e}")

if __name__ == "__main__":
    user_interface()
